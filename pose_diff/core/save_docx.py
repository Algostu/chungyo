from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches
from save_graph import save_graph

def save_docx(video1,video2,docx_name):   #input 값을 지정해주어 원하는 값을 넣으면 원하는 아웃풋이 나오도록 바꾸기
    document=Document()

    save_graph(1,1,video1,video2,'왼쪽 겨드랑이')
    save_graph(3,3, video1, video2, '오른쪽 겨드랑이')
    save_graph(2,2,video1,video2,'왼쪽 팔꿈치')
    save_graph(4,4, video1, video2, '오른쪽 팔꿈치')
    save_graph(1,3,video1,video1,'왼오른 겨드랑이')
    save_graph(2,4,video1,video1,'왼오른 팔꿈치')

    head=document.add_heading('운동 분석(User-Trainer)',0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('겨드랑이 각도(User-RED,Trainer-BLUE)')
    document.add_picture('왼쪽 겨드랑이.png',width=Inches(3))
    document.add_picture('오른쪽 겨드랑이.png', width=Inches(3))
    document.add_paragraph('팔꿈치 각도(User-RED,Trainer-BLUE)')
    document.add_picture('왼쪽 팔꿈치.png',width=Inches(3))
    document.add_picture('오른쪽 팔꿈치.png', width=Inches(3))
    #포멧 또는 % 사용해서 바꾸기
    document.add_page_break()

    document.add_heading('Trainer와 나의 카메라 위치 다르면 값이 다르게 나올 수 있습니다.\n\n', level=2)
    document.add_heading('Trainer의 그래프보다 내 그래프가 앞(뒤)에 있는 경우',level=3)
    document.add_paragraph('Trainer 보다 먼저(늦게) 시작했다\n 시작지점을 조금만 신경 써 주세요!')
    document.add_paragraph('')
    document.add_heading('Trainer의 그래프보다 내 그래프가 위아래로 짧은 경우',level=3)
    document.add_paragraph('Trainer 보다 덜 팔이 구부려진다 또는 덜 움직인다.\n 무게를 낮추거나 운동에 도움을 주는 기구를 사용해주세요.')
    document.add_heading('Trainer의 그래프보다 내 그래프가 위아래로 긴 경우', level=3)
    document.add_paragraph('Trainer 보다 더 가동범위가 넓다.\n 운동에 따라 가동범위가 넓으면 부상의 위험이 있을 수 있으니 확인 후 다시 운동해주세요.')
    document.add_paragraph('부상의 위험이 없다면 매우 잘하고 있는 중 입니다.')
    document.add_heading('Trainer의 그래프보다 내 그래프가 넓은(좁은) 경우',level=3)
    document.add_paragraph('Trainer 보다 속도가 느리(빠르다)\n 속도에 조금 더 신경 써 주세요.')
    document.add_heading('Trainer와 내가 비슷한 경우',level=3)
    document.add_paragraph('잘하는 중이다.\n')
    document.add_page_break()


    document.add_heading('왼손과 오른손의 움직임을 비교하는 그래프이다.\n',level=2)
    document.add_picture('왼오른 겨드랑이.png', width=Inches(5))
    document.add_picture('왼오른 팔꿈치.png', width=Inches(5))
    document.add_page_break()
    document.add_heading('카메라가 정면에 있지 않으면 좌우의 값이 이상하게 나올 수 있습니다..\n\n', level=2)
    document.add_heading('그래프가 위 아래로 다른 경우',level=3)
    document.add_paragraph('좌우의 균형이 안맞습니다. \n  좀 더 낮게 나온 부위를 신경을 쓰면서 같은 힘을 줄 수 있도록 노력하세요 ')
    document.add_heading('그래프가 좌우로 다른 경우', level=3)
    document.add_paragraph('카메라가 정면에 있는지 확인해주세요.')

    document.save(docx_name)




